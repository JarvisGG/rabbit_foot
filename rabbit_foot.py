 #-*- coding:utf-8 -*-
import os
import os.path
import docx
import sys
    
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from sys import argv

document = Document()
total_page_num = 0
page_line_count = 0

def findJavaFile(path):
    global total_page_num
    for parent, dirnames, filenames in os.walk(path):
        for dirname in dirnames:
            print("parent is:" + parent)
            print("dirname is:" + dirname)
        
        for filename in filenames:
            print("the full name of the file is:" + os.path.join(parent,filename))
            if total_page_num >= 80:
                break
            operatorFiles(os.path.join(parent, filename))
            
def operatorFiles(file_path):

    global total_page_num
    global page_line_count
    if file_path.endswith(".java") and (not file_path.endswith('R.java') and (not file_path.endswith('R2.java'))):
        with open(file_path, 'r', encoding ='utf-8') as f:
            for line in f.readlines():
                if len(line) == 0:
                    continue
                page_line_count += 1
                if page_line_count == 50:
                    document.add_page_break()
                    total_page_num += 1
                    page_line_count = 0
                paragraph = document.add_paragraph(line, style='List Number')

def initDocument():
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(4)

if __name__ == '__main__':
    path = ''
    project_name = ''
    for i in range(1, len(argv)):
        if i is 1:
            path = argv[i]
            print(path)
            argvlist = path.split('/')
        elif i is 2:
            project_name = argv[i]
            print(project_name)

    initDocument()
    document.add_heading(project_name, 0)
    findJavaFile(path)
    document.save('./'+project_name+'.docx')